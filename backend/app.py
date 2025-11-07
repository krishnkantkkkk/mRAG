import hashlib
import os
import pickle
import io
import fitz
import faiss
import docx
import tempfile
from werkzeug.utils import secure_filename
from flask import Flask, Response, request, jsonify, stream_with_context, json
from flask_cors import CORS
from sentence_transformers import SentenceTransformer, CrossEncoder
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_ollama.chat_models import ChatOllama
from langchain_text_splitters import RecursiveCharacterTextSplitter
from faster_whisper import WhisperModel
import pytesseract
from PIL import Image
import numpy as np
from concurrent.futures import ThreadPoolExecutor
import threading
from typing import TypedDict, Annotated


app = Flask(__name__, static_folder='temp', static_url_path='/temp')
CORS(app)


all_documents_metadata = []
vector_store = None
session_uploaded_files = set()
session_file_hashes = {}
session_file_indices = {}
CACHE_DIR = "cache"


_embedding_model = None
_reranker = None
_whisper_model = None
_ex_llm = None
_llm = None
_model_lock = threading.Lock()

_blip_processor = None
_blip_model = None

def get_blip_model():
    global _blip_processor, _blip_model
    if _blip_model is None:
        with _model_lock:
            if _blip_model is None:
                from transformers import BlipProcessor, BlipForConditionalGeneration
                _blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-large")
                _blip_model = BlipForConditionalGeneration.from_pretrained(
                    "Salesforce/blip-image-captioning-large"
                ).to("cpu")
    return _blip_processor, _blip_model


def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        with _model_lock:
            if _embedding_model is None:
                _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    return _embedding_model


def get_reranker():
    global _reranker
    if _reranker is None:
        with _model_lock:
            if _reranker is None:
                _reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
    return _reranker


def get_whisper_model():
    global _whisper_model
    if _whisper_model is None:
        with _model_lock:
            if _whisper_model is None:
                _whisper_model = WhisperModel("base.en", device="cpu", compute_type="int8")
    return _whisper_model


def get_llm():
    global _llm
    if _llm is None:
        with _model_lock:
            if _llm is None:
                _llm = ChatOllama(model="llama3.2")
    return _llm


def get_ex_llm():
    global _ex_llm
    if _ex_llm is None:
        with _model_lock:
            if _ex_llm is None:
                _ex_llm = ChatOllama(model="gemma3:1b")
    return _ex_llm


executor = ThreadPoolExecutor(max_workers=4)


def describe_image_with_vision_model(image_path, context_before="", context_after=""):
    try:
        processor, model = get_blip_model()
        image = Image.open(image_path).convert("RGB")
        inputs = processor(image, return_tensors="pt")

        outputs = model.generate(
            **inputs,
            max_length=100,
            num_return_sequences=5,
            do_sample=True,
            top_k=50,
            top_p=0.95
        )
        caption = "Here are some descriptions of the image: "
        for out in outputs:
            caption += processor.decode(out, skip_special_tokens=True) + " "
        print(caption)
        image_caption_llm = get_ex_llm()
        caption_prompt = ChatPromptTemplate.from_template(f"""Based on the following caption and surrounding text context,\n                                    
        Context Before Image: [{context_before}]\n
        Caption: [{caption}]\n
        Context After Image: [{context_after}]\n
        Provide a concise and relevant description of the image in three sentences.
        what can be the type of image [graph, chart, diagram, potrait or photograph] 
        If potrait/photograph who can be in the image what could be the name.""")
        caption_chain = caption_prompt | image_caption_llm | StrOutputParser()
        caption = caption_chain.invoke({})
        caption = f"Image Description: [{caption.strip()}]"
        print(caption)
        return caption
    except Exception as e:
        print(f"Vision model description failed: {e}")
        return "Unable to generate detailed image description."


def process_pdf(file_storage):
    file_bytes = io.BytesIO(file_storage.read())
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    
    processed_data = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)

    def process_page(page_num):
        page = doc[page_num]
        page_data = []
        
        # Process text
        text = page.get_text()
        if text.strip():
            chunks = text_splitter.split_text(text)
            for chunk in chunks:
                page_data.append({
                    "text": chunk,
                    "page_num": page_num + 1,
                    "source_filename": file_storage.filename,
                    "type": "text"
                })

        # Process images with vision model
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                pil_image = Image.open(io.BytesIO(img_bytes)).convert("RGB")

                os.makedirs("temp", exist_ok=True)
                ext = base_image["ext"]
                img_filename = f"{os.path.splitext(file_storage.filename)[0]}_p{page_num+1}_{img_index}.{ext}"
                img_path = os.path.join("temp", img_filename)
                pil_image.save(img_path)

                ocr_text = pytesseract.image_to_string(pil_image)
                vision_description = describe_image_with_vision_model(img_path)
                
                image_doc_text = f"""Image from page {page_num + 1} of {file_storage.filename}
Vision Description: {vision_description}
OCR Text: {ocr_text.strip()}""".strip()

                page_data.append({
                    "text": image_doc_text,
                    "image_path": img_filename,
                    "page_num": page_num + 1,
                    "source_filename": file_storage.filename,
                    "type": "image",
                })
            except Exception as e:
                print(f"Warning: Could not process image {img_index} on page {page_num+1}: {e}")
        
        return page_data
    
    with ThreadPoolExecutor(max_workers=4) as page_executor:
        page_results = list(page_executor.map(process_page, range(len(doc))))
    
    for page_data in page_results:
        processed_data.extend(page_data)

    doc.close()
    return processed_data

def process_docx(file_storage):

    file_bytes = io.BytesIO(file_storage.read())
    doc = docx.Document(file_bytes)
    
    doc_data = []
    os.makedirs("temp", exist_ok=True)
    
    image_rels = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_rels[rel.rId] = rel
    
    position = 0
    current_text = ""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    
    all_elements = []
    for element in doc.element.body:
        if element.tag.endswith('p'):

            has_image = False
            for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                for drawing in run.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                    blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if blip is not None:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in image_rels:

                            para_text = ""
                            for para in doc.paragraphs:
                                if para._element == element:
                                    para_text = para.text
                                    break
                            if para_text.strip():
                                all_elements.append({"type": "text", "content": para_text})
                            
                            all_elements.append({"type": "image", "embed": embed, "rel": image_rels[embed]})
                            has_image = True
            
            if not has_image:
                para_text = ""
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text
                        break
                
                if para_text.strip():
                    all_elements.append({"type": "text", "content": para_text})
    
    for idx, elem_info in enumerate(all_elements):
        if elem_info["type"] == "text":
            para_text = elem_info["content"]
            current_text += para_text + "\n"
            
            if len(current_text) >= 1000:
                chunks = text_splitter.split_text(current_text.strip())
                for chunk in chunks:
                    position += 1
                    doc_data.append({
                        "text": chunk,
                        "page_num": position,
                        "source_filename": file_storage.filename,
                        "type": "text"
                    })
                current_text = ""
        
        elif elem_info["type"] == "image":
            context_before = ""
            if current_text.strip():
                chunks = text_splitter.split_text(current_text.strip())
                for chunk in chunks:
                    position += 1
                    doc_data.append({
                        "text": chunk,
                        "page_num": position,
                        "source_filename": file_storage.filename,
                        "type": "text"
                    })
                context_before = current_text[-500:] if len(current_text) > 500 else current_text
                current_text = ""
            else:
                for prev_doc in reversed(doc_data[-3:]):
                    if prev_doc['type'] == 'text':
                        context_before = prev_doc['text'][:300] + " " + context_before
                context_before = context_before.strip()[:500]
            
            context_after = ""
            for next_idx in range(idx + 1, min(idx + 4, len(all_elements))):
                if all_elements[next_idx]["type"] == "text":
                    context_after += all_elements[next_idx]["content"] + " "
                    if len(context_after) >= 500:
                        break
            context_after = context_after.strip()[:500]
            
            try:
                rel = elem_info["rel"]
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data)).convert("RGB")
                
                ext = rel.target_ref.split('.')[-1]
                img_filename = f"{os.path.splitext(file_storage.filename)[0]}_img{position}.{ext}"
                img_path = os.path.join("temp", img_filename)
                img.save(img_path)
                
                ocr_text = pytesseract.image_to_string(img)
                
                vision_description = describe_image_with_vision_model(img_path, context_before, context_after)
                
                image_doc_text = f"""Image from {file_storage.filename}
Vision Description: {vision_description}
OCR Text: {ocr_text.strip()}""".strip()
                
                position += 1
                doc_data.append({
                    "text": image_doc_text,
                    "image_path": img_filename,
                    "page_num": position,
                    "source_filename": file_storage.filename,
                    "type": "image",
                })
            except Exception as e:
                print(f"Warning: Could not process DOCX image at position {position}: {e}")
    
    if current_text.strip():
        chunks = text_splitter.split_text(current_text.strip())
        for chunk in chunks:
            position += 1
            doc_data.append({
                "text": chunk,
                "page_num": position,
                "source_filename": file_storage.filename,
                "type": "text"
            })
    
    return doc_data

def process_audio(file_storage):
    import re
    
    os.makedirs("temp", exist_ok=True)
    filename = secure_filename(file_storage.filename)
    audio_path = os.path.join("temp", filename)
    file_storage.save(audio_path)

    whisper = get_whisper_model()
    segments, _ = whisper.transcribe(audio_path, word_timestamps=True)

    # Parse segments into timestamped sentences
    timestamped_segments = []
    
    for segment in segments:
        if not segment.words:
            continue
            
        current_sentence = {
            "text": "",
            "start_time": None,
            "end_time": None
        }
        
        for word in segment.words:
            # Initialize start time
            if current_sentence["start_time"] is None:
                current_sentence["start_time"] = word.start
            
            # Add word to sentence
            current_sentence["text"] += word.word + " "
            current_sentence["end_time"] = word.end
            
            # Check if sentence ends (., ?, !)
            if word.word.strip() and word.word.strip()[-1] in ['.', '?', '!']:
                # Save the completed sentence
                if current_sentence["text"].strip():
                    timestamped_segments.append({
                        "text": current_sentence["text"].strip(),
                        "start_time": current_sentence["start_time"],
                        "end_time": current_sentence["end_time"]
                    })
                
                # Reset for next sentence
                current_sentence = {
                    "text": "",
                    "start_time": None,
                    "end_time": None
                }
        
        # Handle remaining text (sentence without punctuation)
        if current_sentence["text"].strip():
            timestamped_segments.append({
                "text": current_sentence["text"].strip(),
                "start_time": current_sentence["start_time"],
                "end_time": current_sentence["end_time"]
            })
    
    if not timestamped_segments:
        return []
    
    # Now create semantic chunks from sentences
    chunk_data = []
    max_chunk_size = 1000  # characters
    overlap_sentences = 2  # number of sentences to overlap
    
    i = 0
    chunk_counter = 0
    
    while i < len(timestamped_segments):
        current_chunk_text = ""
        current_chunk_sentences = []
        chunk_start_idx = i
        
        # Build chunk by adding sentences until we exceed max_chunk_size
        while i < len(timestamped_segments):
            sentence = timestamped_segments[i]
            
            # Check if adding this sentence would exceed limit
            if current_chunk_text and len(current_chunk_text + " " + sentence["text"]) > max_chunk_size:
                break
            
            # Add sentence to chunk
            if current_chunk_text:
                current_chunk_text += " " + sentence["text"]
            else:
                current_chunk_text = sentence["text"]
            
            current_chunk_sentences.append(sentence)
            i += 1
        
        # Create chunk with overall start/end times
        if current_chunk_sentences:
            chunk_counter += 1
            chunk_start_time = current_chunk_sentences[0]["start_time"]
            chunk_end_time = current_chunk_sentences[-1]["end_time"]
            
            # Build detailed text with inline timestamps for each sentence
            detailed_text_parts = []
            for sent in current_chunk_sentences:
                sent_with_time = f"[{sent['start_time']:.2f}s - {sent['end_time']:.2f}s] {sent['text']}"
                detailed_text_parts.append(sent_with_time)
            
            detailed_text = "\n".join(detailed_text_parts)
            
            chunk_data.append({
                "text": detailed_text,
                "page_num": chunk_counter,
                "source_filename": file_storage.filename,
                "type": "audio",
                "start_time": chunk_start_time,
                "end_time": chunk_end_time,
                "duration": chunk_end_time - chunk_start_time,
                "sentence_count": len(current_chunk_sentences)
            })
        
        # Handle overlap: backtrack by overlap_sentences
        if i < len(timestamped_segments) and len(current_chunk_sentences) > overlap_sentences:
            i = i - overlap_sentences
    
    return chunk_data

def create_vector_store_from_docs(documents):
    embedding_model = get_embedding_model()
    texts = [doc['text'] for doc in documents]
    embeddings = embedding_model.encode(texts, convert_to_tensor=True, show_progress_bar=False)
    
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings.cpu().numpy())
    return index
    

def format_timestamp(seconds):
    if seconds is None:
        return "00:00"
    total = int(round(seconds))
    minutes, secs = divmod(total, 60)
    return f"{minutes:02d}:{secs:02d}"


def expand_query(query):
#     print("Expanding query for better retrieval...")
#     template = """Based on the user's question, generate 3 additional, different, and more specific queries that are likely to find relevant documents in a vector database.
# Focus on rephrasing, using synonyms, and breaking down the question into sub-questions.
# Provide ONLY the queries, each on a new line. Do not number them or add any other text. And just questions no explainations, nothing else.

# Original Question: {question}

# Generated Queries:"""
    
#     llm = get_ex_llm()
#     prompt = ChatPromptTemplate.from_template(template)
#     chain = prompt | llm | StrOutputParser()
    
#     try:
#         response = chain.invoke({"question": query})
#         expanded_queries = [q.strip() for q in response.strip().split('\n') if q.strip()]
#         all_queries = [query] + expanded_queries[:3]
#         print("Expanded Queries:", all_queries)
#         return list(set(all_queries))
#     except Exception as e:
#         print(f"Query expansion failed: {e}")
        return [query]


def load_from_cache(file_hash):
    cache_path = os.path.join(CACHE_DIR, file_hash)
    docs_path = os.path.join(cache_path, "documents.pkl")
    embeddings_path = os.path.join(cache_path, "embeddings.npy")
    
    if os.path.exists(docs_path):
        try:
            with open(docs_path, "rb") as f:
                docs = pickle.load(f)
            
            embeddings = None
            if os.path.exists(embeddings_path):
                embeddings = np.load(embeddings_path)
                
            return {
                "docs": docs,
                "embeddings": embeddings
            }
        except Exception as e:
            print(f"Could not load cache for {file_hash}: {e}")
    return None


def save_to_cache(file_hash, docs, embeddings=None):
    """Save processed document data and embeddings to cache."""
    cache_path = os.path.join(CACHE_DIR, file_hash)
    os.makedirs(cache_path, exist_ok=True)
    
    with open(os.path.join(cache_path, "documents.pkl"), "wb") as f:
        pickle.dump(docs, f)
    
    if embeddings is not None:
        np.save(os.path.join(cache_path, "embeddings.npy"), embeddings)


@app.route('/upload', methods=['POST'])
def upload_file():
    global all_documents_metadata, vector_store, session_uploaded_files, session_file_hashes, session_file_indices
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        newly_processed_docs = []
        processed_filenames = []
        
        all_new_embeddings = []
        all_new_docs = []

        for file in files:
            file_bytes = file.read()
            file.seek(0)
            file_hash = hashlib.sha256(file_bytes).hexdigest()
            
            if file.filename in session_uploaded_files and session_file_hashes.get(file.filename) == file_hash:
                print(f"File {file.filename} already uploaded, skipping...")
                continue
            
            cached_data = load_from_cache(file_hash)
            
            if cached_data is not None and cached_data["docs"] is not None:
                print(f"Loading {file.filename} from cache...")
                docs = cached_data["docs"]
                cached_embeddings = cached_data["embeddings"]
                
                for doc in docs:
                    doc['source_filename'] = file.filename
                
                if cached_embeddings is not None:
                    print(f"Using cached embeddings for {file.filename}")
                    new_embeddings = cached_embeddings
                else:
                    print(f"Creating embeddings for cached documents of {file.filename}")
                    embedding_model = get_embedding_model()
                    new_texts = [doc['text'] for doc in docs]
                    new_embeddings = embedding_model.encode(new_texts, convert_to_tensor=True, show_progress_bar=False).cpu().numpy()
                    save_to_cache(file_hash, docs, new_embeddings)
                    
            else:
                print(f"Processing new file: {file.filename}")
                
                filename = file.filename.lower()
                if filename.endswith('.pdf'):
                    docs = process_pdf(file)
                elif filename.endswith('.docx'):
                    docs = process_docx(file)
                elif filename.endswith(('.mp3', '.wav', '.m4a', '.ogg')):
                    docs = process_audio(file)
                else:
                    continue
                
                if not docs: 
                    continue
                
                new_texts = [doc['text'] for doc in docs]
                embedding_model = get_embedding_model()
                new_embeddings = embedding_model.encode(new_texts, convert_to_tensor=True, show_progress_bar=False).cpu().numpy()
                
                save_to_cache(file_hash, docs, new_embeddings)
            
            start_idx = len(all_documents_metadata) + len(all_new_docs)
            end_idx = start_idx + len(docs)
            session_file_indices[file.filename] = {
                "start": start_idx,
                "end": end_idx,
                "count": len(docs)
            }
            
            all_new_docs.extend(docs)
            all_new_embeddings.append(new_embeddings)
            session_uploaded_files.add(file.filename)
            session_file_hashes[file.filename] = file_hash
            processed_filenames.append(file.filename)

        if all_new_embeddings:
            combined_embeddings = np.vstack(all_new_embeddings)
            
            if vector_store is None:
                dimension = combined_embeddings.shape[1]
                vector_store = faiss.IndexFlatL2(dimension)
            
            vector_store.add(combined_embeddings)
            all_documents_metadata.extend(all_new_docs)
            
            print(f"Added {len(all_new_docs)} total chunks from {len(processed_filenames)} files")

        return jsonify({'message': 'Files processed successfully', 'filenames': processed_filenames})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/ask', methods=['POST'])
def ask_question():
    if vector_store is None or vector_store.ntotal == 0: 
        return jsonify({'error': 'No documents uploaded yet'}), 400
    
    data = request.get_json()
    question = data.get('question')
    if not question: 
        return jsonify({'error': 'No question provided'}), 400

    queries = expand_query(question)
    
    embedding_model = get_embedding_model()
    query_embeddings = embedding_model.encode(queries, convert_to_tensor=True, show_progress_bar=False).cpu().numpy()

    k_retrieval = 15
    distances, ids = vector_store.search(query_embeddings, k_retrieval)
    
    unique_ids = set()
    for id_list in ids:
        for i in id_list:
            if i != -1 and 0 <= i < len(all_documents_metadata):
                unique_ids.add(i)
    
    candidate_docs = [all_documents_metadata[i] for i in unique_ids]
    
    if not candidate_docs:
        return Response(stream_with_context(iter(["<div><p>I couldn't find any relevant information in the uploaded documents to answer your question.</p></div>"])))

    reranker = get_reranker()
    rerank_pairs = [[question, doc['text']] for doc in candidate_docs]
    scores = reranker.predict(rerank_pairs)
    
    doc_scores = list(zip(candidate_docs, scores))
    doc_scores.sort(key=lambda x: x[1], reverse=True)
    
    # Check if query is about images/visuals or audio
    is_visual_query = any(keyword in question.lower() for keyword in ['image', 'graph', 'chart', 'diagram', 'picture', 'show'])
    is_audio_query = any(keyword in question.lower() for keyword in ['audio', 'said', 'mentioned', 'timestamp', 'when', 'time'])
    has_audio_docs = any(doc.get('type') == 'audio' for doc in candidate_docs)
    
    # Adjust top_k based on query type
    top_k_reranked = 2
    retrieved_results = doc_scores[:top_k_reranked]
    retrieved_docs_metadata = [res[0] for res in retrieved_results]
    
    # Build context with timestamp information for audio chunks
    context_parts = []
    for doc in retrieved_docs_metadata:
        context_part = f"Source from {doc['source_filename']}, Chunk {doc.get('page_num', 'N/A')}:\n"
        
        # For audio, the text already contains inline timestamps per sentence
        # Just add the overall chunk time range
        if doc.get('type') == 'audio' and doc.get('start_time') is not None:
            context_part += f"[Overall chunk time range: {doc.get('start_time'):.2f}s - {doc.get('end_time'):.2f}s]\n\n"
        
        context_part += doc['text']
        context_parts.append(context_part)
    
    context_text = "\n\n".join(context_parts)

    # Check if any retrieved documents are audio
    contains_audio = any(doc.get('type') == 'audio' for doc in retrieved_docs_metadata)

    class LLMResponse(TypedDict):
        m_s_g: Annotated[str, "The answer ONLY from the context in plain text or HTML <div>...</div> format with formatting."]
        f_n_d: Annotated[bool, "True or False, Relevant Answer found or not."]
        st_time: Annotated[float | None, "Start time in seconds if context have [23.33 - 33.44] format return most accurate start time"]
        en_time: Annotated[float | None, "End time in seconds if context have [34.33 - 78.44] format return most accurate end time"]

    prompt_template = """
Answer ONLY based on the given context.
If asked for showing image you have to just provide one line caption if available.

Context: {context}

Question: {question}
Provide your answer in the following JSON format:
{{  
"m_s_g": "<Your answer here in plain text or HTML <div>...</div> format and use html tags when needed.>",
"f_n_d": <True or False, Relevant Answer found or not.>,
"st_time": <Start time in seconds if context have [23.33 - 33.44] format return most accurate and precise start time or null>,'
"en_time": <End time in seconds if context have [34.33 - 78.44] format return most accurate and precise end time or null>
}}
"""

    prompt = ChatPromptTemplate.from_template(prompt_template)

    llm = get_llm()
    rag_chain = prompt | llm.with_structured_output(LLMResponse)

    def generate():
        full_response = {"message": "", "found": False, "start_time": None, "end_time": None}
        found = False
        response_complete = False
        start_time = None
        end_time = None

        try:
            for chunk in rag_chain.stream({"context": context_text, "question": question}):
                if isinstance(chunk, dict):
                    if "m_s_g" in chunk:
                        msg = chunk["m_s_g"]
                        new_text = msg[len(full_response['message']):]
                        full_response['message'] = msg
                        if new_text:
                            yield new_text
                    
                    if "f_n_d" in chunk:
                        full_response['found'] = chunk["f_n_d"]
                        found = chunk["f_n_d"]
                    if "st_time" in chunk:
                        full_response['start_time'] = chunk["st_time"]
                    if "en_time" in chunk:
                        full_response['end_time'] = chunk["en_time"]

            print(full_response)
                        
        except Exception as e:
            print(f"Error during streaming: {e}")
            import traceback
            traceback.print_exc()
            yield f"\n\n[Error processing response: {str(e)}]"
            return

        # Only show sources if answer was found
        if not found:
            return

        # Determine if images should be shown inline
        should_show_images = any(keyword in question.lower() for keyword in ['show', 'display', 'image', 'graph', 'chart', 'diagram', 'picture'])
        
        # Prepare sources for display
        sources = []
        for doc, score in retrieved_results:
            source_obj = {
                "source_filename": doc['source_filename'],
                "page_num": doc.get('page_num', 0),
                "source_content": doc['text'][:500] + "..." if len(doc['text']) > 500 else doc['text'],
                "type": doc.get('type', 'text'),
                "score": float(score)
            }
            
            # Add image-specific fields
            if doc.get('type') == 'image': 
                source_obj['image_path'] = doc.get('image_path', '')
                source_obj['vision_description'] = doc.get('vision_description', '')
                source_obj['show_inline'] = should_show_images
            
            # Add audio-specific fields
            if doc.get('type') == 'audio':
                source_obj['start_time'] = full_response["start_time"]
                source_obj['end_time'] = full_response["end_time"]
                source_obj['duration'] = full_response["end_time"] - full_response["start_time"] if full_response["start_time"] is not None and full_response["end_time"] is not None else None
                source_obj['timestamp_display'] = f"{format_timestamp(full_response['start_time'])} - {format_timestamp(full_response['end_time'])}"
            
            sources.append(source_obj)
        
        # Add LLM-identified precise timestamps if available
        source_metadata = {
            "type": "sources",
            "content": sources
        }
        
        if start_time is not None and end_time is not None:
            source_metadata["precise_timestamp"] = {
                "start_time": start_time,
                "end_time": end_time,
                "timestamp_display": f"{format_timestamp(start_time)} - {format_timestamp(end_time)}"
            }
        
        # Send sources as JSON
        yield "\n\n" + json.dumps(source_metadata)

    return Response(stream_with_context(generate()), mimetype='text/plain')

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio:
            audio_file.save(temp_audio.name)
            temp_audio_path = temp_audio.name
        
        try:
            whisper = get_whisper_model()
            segments, _ = whisper.transcribe(temp_audio_path, beam_size=5)
            transcription = " ".join([segment.text for segment in segments])
        finally:
            os.remove(temp_audio_path)
            
        return jsonify({'transcription': transcription})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/files', methods=['GET'])
def list_files():
    file_info = []
    for filename in sorted(session_uploaded_files):
        info = {
            "filename": filename,
            "hash": session_file_hashes.get(filename),
            "indices": session_file_indices.get(filename, {}),
            "chunk_count": session_file_indices.get(filename, {}).get("count", 0)
        }
        file_info.append(info)
    
    return jsonify({
        'files': list(session_uploaded_files),
        'detailed_info': file_info,
        'total_chunks': len(all_documents_metadata),
        'vector_store_size': vector_store.ntotal if vector_store else 0
    })


@app.route('/session-info', methods=['GET'])
def session_info():
    return jsonify({
        'uploaded_files': list(session_uploaded_files),
        'file_indices': session_file_indices,
        'total_documents': len(all_documents_metadata),
        'vector_store_size': vector_store.ntotal if vector_store else 0,
        'cache_stats': {
            'cached_files': len([d for d in os.listdir(CACHE_DIR) if os.path.isdir(os.path.join(CACHE_DIR, d))]) if os.path.exists(CACHE_DIR) else 0
        }
    })


@app.route('/clear-session', methods=['POST'])
def clear_session():
    global all_documents_metadata, vector_store, session_uploaded_files, session_file_hashes, session_file_indices
    
    all_documents_metadata = []
    vector_store = None
    session_uploaded_files.clear()
    session_file_hashes.clear()
    session_file_indices.clear()
    
    return jsonify({'message': 'Session cleared successfully'})

if __name__ == '__main__':
    app.run(debug=True, port=5000, threaded=True)